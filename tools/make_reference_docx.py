"""Build the submission reference DOCX used by the DOCX renderer.

Run this once (or whenever the manuscript house style changes) to regenerate
``src/epy_paper/assets/reference_docx/submission.docx``. Pandoc copies the
paragraph and character styles from this reference document into every DOCX
draft, so the manuscript inherits the universal submission look:

- US Letter, 1-inch margins.
- Times New Roman 12 pt body.
- Double line spacing.
- Heading styles Pandoc maps to (``Title``, ``Heading 1..4``).

This is a build-time tool, not a runtime dependency: it requires
``python-docx`` (installed via the ``dev`` extra), while the library itself
only needs ``pypandoc``.

Usage::

    python tools/make_reference_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)

_OUT = (
    Path(__file__).resolve().parent.parent
    / "src"
    / "epy_paper"
    / "assets"
    / "reference_docx"
    / "submission.docx"
)


def _set_double_spacing(paragraph_format) -> None:
    """Apply double line spacing to a paragraph format."""
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE


def build() -> Path:
    """Build the reference DOCX and return its path."""
    doc = Document()

    # Page geometry: US Letter, 1-inch margins.
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in (
        "top_margin",
        "bottom_margin",
        "left_margin",
        "right_margin",
    ):
        setattr(section, attr, Inches(1))

    # Normal style: Times New Roman 12 pt, double spaced.
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    _set_double_spacing(normal.paragraph_format)

    # Headings inherit the body font; keep them readable and serif.
    for name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = BODY_FONT
        _set_double_spacing(style.paragraph_format)

    # Seed paragraphs so Pandoc has every style instantiated to copy.
    doc.add_paragraph("Title", style="Title")
    for level in range(1, 5):
        doc.add_paragraph(f"Heading {level}", style=f"Heading {level}")
    doc.add_paragraph(
        "Body text in Times New Roman 12 pt, double spaced.",
        style="Normal",
    )

    _OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(_OUT))
    return _OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
