"""Build the submission reference DOCX used by the DOCX renderer.

Run this once (or whenever the manuscript house style changes) to regenerate
``src/epy_papers/_config/_assets/reference_docx/submission.docx``. Pandoc
copies the paragraph and character styles from this reference document into
every DOCX draft, so the manuscript inherits the universal submission look:

- US Letter, 1-inch margins.
- Times New Roman 12 pt body.
- Double line spacing.
- Heading styles Pandoc maps to (``Title``, ``Heading 1..4``).

This is a build-time tool, not a runtime dependency: it requires
``python-docx`` (installed via the ``dev`` extra), while the library itself
only needs ``pypandoc``.

Usage::

    python src/epy_papers/_core/_packaging/tools/make_reference_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)

_REF_DIR = (
    Path(__file__).resolve().parents[5]
    / "src"
    / "epy_papers"
    / "_config"
    / "_assets"
    / "reference_docx"
)


def _add_line_numbering(section) -> None:
    """Enable continuous line numbering on a section's properties.

    Pandoc copies the reference document's section properties into the draft,
    so journals whose profile asks for numbered lines get them in the DOCX.
    """
    sect_pr = section._sectPr
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "continuous")
    ln.set(qn("w:distance"), "360")
    cols = sect_pr.find(qn("w:cols"))
    if cols is not None:
        cols.addprevious(ln)
    else:
        sect_pr.append(ln)


def _set_double_spacing(paragraph_format) -> None:
    """Apply double line spacing to a paragraph format."""
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE


def build(line_numbered: bool = False) -> Path:
    """Build a reference DOCX and return its path.

    With ``line_numbered`` the section enables continuous line numbering so
    Pandoc-produced drafts inherit it for journals that require numbered lines.
    """
    doc = Document()

    # Page geometry: US Letter, 1-inch margins.
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in (
        "top_margin",
        "bottom_margin",
        "left_margin",
        "right_margin",
    ):
        setattr(section, attr, Inches(1))

    # Normal style: Times New Roman 12 pt, double spaced.
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    _set_double_spacing(normal.paragraph_format)

    # Headings inherit the body font; keep them readable and serif.
    for name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = BODY_FONT
        _set_double_spacing(style.paragraph_format)

    # Seed paragraphs so Pandoc has every style instantiated to copy.
    doc.add_paragraph("Title", style="Title")
    for level in range(1, 5):
        doc.add_paragraph(f"Heading {level}", style=f"Heading {level}")
    doc.add_paragraph(
        "Body text in Times New Roman 12 pt, double spaced.",
        style="Normal",
    )

    if line_numbered:
        _add_line_numbering(section)

    out = _REF_DIR / (
        "submission_lineno.docx" if line_numbered else "submission.docx"
    )
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


if __name__ == "__main__":
    for _line_numbered in (False, True):
        print(f"Wrote {build(_line_numbered)}")
